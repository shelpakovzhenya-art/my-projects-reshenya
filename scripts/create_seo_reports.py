from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path.home() / "Downloads"
TODAY = "22.06.2026"


GUIDELINES = [
    "Яндекс.Вебмастер: канонический адрес страницы, дубли, robots.txt, sitemap.xml, ответы сервера.",
    "Google Search Central: canonical для дублей, robots.txt, sitemap, редиректы, page metadata.",
    "Материалы из таблицы знаний: SEO-курсы, Яндекс-сертификация, Rush Academy, Senior/Middle SEO, материалы SEO-сообществ.",
]


REPORTS = [
    {
        "project": "Фильтр-пресс",
        "url": "https://filter-press.ru/",
        "access": "Bitrix-админка доступна, вход выполнен. Yandex Webmaster не проверен: браузер остановился на форме входа Яндекс ID.",
        "summary": "Критичная проблема — отсутствие склейки зеркал: http, https, www и non-www отдают 200 OK. Дополнительно доступен дубль /index.php, а canonical на проверенных страницах не найден.",
        "fixed": [
            "Вошел в Bitrix и проверил наличие штатных SEO-разделов: robots.txt, sitemap.xml, Aspro sitemap, Умный SEO.",
            "Публичный sitemap.xml актуальный по Last-Modified: 16.04.2026, поэтому регенерацию не выполнял.",
            "Изменений не вносил: проблема зеркал и /index.php требует серверных 301-редиректов, а не контентной правки в CMS.",
        ],
        "issues": [
            ("P1", "Склеить зеркала", "Настроить 301: http -> https, www -> non-www. Канонический хост: https://filter-press.ru/."),
            ("P1", "Закрыть /index.php", "Сделать 301 с /index.php и внутренних index.php на соответствующий URL без index.php."),
            ("P1", "Добавить canonical", "Включить rel=canonical на главной, разделах, карточках и посадочных страницах."),
            ("P2", "Параметры URL", "UTM/служебные параметры должны иметь canonical на чистую страницу или редирект."),
        ],
        "evidence": [
            "https://filter-press.ru/, http://filter-press.ru/, https://www.filter-press.ru/ отдают 200 OK.",
            "https://filter-press.ru/index.php отдает 200 OK.",
            "CMS: 1C-Bitrix, шаблон Aspro Allcorp3.",
        ],
    },
    {
        "project": "Грани",
        "url": "https://gs-best.ru/",
        "access": "ICMS-доступ рабочий, вход выполнен. Yandex Webmaster не проверен: требуется вход в Яндекс ID.",
        "summary": "Основная проблема — редиректные цепочки через HTTP и :443, а также отсутствие canonical. Проверенная ранее жалоба на дубли SEO у /o-kompanii не подтвердилась при контрольной проверке: в CMS и на сайте title/description уже уникальные.",
        "fixed": [
            "Вошел в ICMS, проверил структуру и SEO-поля страницы /o-kompanii.",
            "Правку /o-kompanii не выполнял: публичная страница уже отдает уникальный title и description.",
            "Раздел редиректов ICMS проверен: он управляет внутренними путями, а проблема http/www/:443 возникает до приложения, на сервере.",
        ],
        "issues": [
            ("P1", "Исправить редиректные цепочки", "Настроить сервер: сразу вести на https://gs-best.ru/... без www, без HTTP-шага и без :443 в Location."),
            ("P1", "Добавить canonical", "В шаблоне сайта вывести rel=canonical для чистого URL страницы."),
            ("P2", "Параметры URL", "Для UTM/sort/прочих параметров нужен canonical на чистый URL или редирект."),
            ("P2", "Проверить sitemap", "В sitemap должны оставаться только финальные URL без редиректов."),
        ],
        "evidence": [
            "https://gs-best.ru/o-kompanii отдает title: О компании | Студия природного камня «Грани» | Грани.",
            "В ICMS у /o-kompanii заполнены seo_title и seo_description.",
            "Редиректы вида /o-kompanii/ проходят через HTTP и https://gs-best.ru:443/.",
        ],
    },
    {
        "project": "Аквабурмастер",
        "url": "https://aquaburmaster.ru/",
        "access": "ICMS-доступ рабочий, вход выполнен. Вебмастер не проверен без входа Яндекс ID. По постановке favicon игнорировался.",
        "summary": "Критичная проблема — URL со слэшем редиректятся на главную: /uslugi/ -> / и /kontakty/ -> /. Это может склеивать важные страницы с главной. Также отсутствует canonical, а публичный sitemap.xml остался старым файлом 2018 года.",
        "fixed": [
            "Вошел в ICMS v3.5.46 и открыл SEO-раздел.",
            "Запустил штатную генерацию sitemap.xml; CMS вернула сообщение: Файл sitemap.xml сгенерирован.",
            "Контроль показал, что публичный sitemap.xml не обновился: в нем остался комментарий mysitemapgenerator.com и lastmod 2018. Исправление не засчитано как выполненное.",
            "Проверил список SEO-редиректов: правил /uslugi/ -> / и /kontakty/ -> / нет, значит проблема в роутере/сервере.",
        ],
        "issues": [
            ("P1", "Исправить редиректы со слэшем", "Сделать /uslugi/ -> /uslugi и /kontakty/ -> /kontakty либо выбрать единый формат URL без редиректа на главную."),
            ("P1", "Добавить canonical", "Вывести canonical на главной, разделах, статьях и услугах."),
            ("P2", "Починить генерацию sitemap", "Проверить права записи и путь генератора: публичный sitemap.xml не меняется после генерации в CMS."),
            ("P2", "Обновить PHP", "Сайт раскрывает X-Powered-By: PHP/5.6.40; версия устаревшая и небезопасная."),
        ],
        "evidence": [
            "https://aquaburmaster.ru/uslugi/ -> 301 на https://aquaburmaster.ru/.",
            "https://aquaburmaster.ru/kontakty/ -> 301 на https://aquaburmaster.ru/.",
            "sitemap.xml: Last-Modified 26.09.2018, created with mysitemapgenerator.com.",
        ],
    },
    {
        "project": "Московская застава",
        "url": "https://mzastava.ru/",
        "access": "Bitrix-админка найдена, но логин из таблицы не сработал: форма вернула Неверный логин или пароль. Yandex Webmaster не проверен без входа Яндекс ID.",
        "summary": "Основные проблемы: редиректы без слэша уходят через HTTP, /index.php отдает 200 OK, canonical отсутствует, sitemap.xml устарел с 2022 года. Из-за недействующего доступа правки не выполнялись.",
        "fixed": [
            "Проверил доступ к /bitrix/admin/: форма доступна.",
            "Один раз попробовал вход по данным из таблицы; получил ошибку Неверный логин или пароль.",
            "Повторные попытки не выполнял, чтобы не спровоцировать блокировку.",
        ],
        "issues": [
            ("P1", "Исправить HTTP в редиректах", "https://mzastava.ru/offers должен сразу вести на https://mzastava.ru/offers/, без промежуточного http."),
            ("P1", "Закрыть /index.php", "Настроить 301 с /index.php на /."),
            ("P1", "Добавить canonical", "Вывести canonical на главной и типовых страницах."),
            ("P2", "Перегенерировать sitemap", "Текущий Last-Modified: 03.08.2022. Нужна пересборка из Bitrix."),
        ],
        "evidence": [
            "CMS: 1C-Bitrix.",
            "sitemap.xml: Last-Modified 03.08.2022.",
            "https://mzastava.ru/index.php отдает 200 OK.",
        ],
    },
    {
        "project": "Красная мельница",
        "url": "https://www.kmel.ru/",
        "access": "ICMS-доступ рабочий, вход выполнен. Yandex Webmaster не проверен: требуется вход в Яндекс ID.",
        "summary": "Ключевые технические проблемы: параметры на валидных страницах дают 404, canonical отсутствует, слэш/без слэша местами редиректится 302, /index.php идет через лишний HTTP-шаг. Безопасная контентная правка выполнена: обновлен meta description главной.",
        "fixed": [
            "Вошел в ICMS v2.2.",
            "В редакторе главной страницы заменил переспамленный meta description на естественное описание.",
            "Публичная проверка подтвердила новый meta description на главной.",
        ],
        "issues": [
            ("P1", "Параметры не должны давать 404", "UTM и служебные параметры должны отдавать чистую страницу с canonical или редиректиться на чистый URL."),
            ("P1", "Добавить canonical", "Вывести canonical на главной, /about_us, /catalog, /catalog/1-4, /contacts."),
            ("P2", "302 заменить на 301", "Слэш/без слэша должен редиректиться постоянным 301."),
            ("P2", "Убрать лишние HTTP-цепочки", "/index.php должен сразу вести на https://www.kmel.ru/."),
        ],
        "evidence": [
            "Измененный description: Красная мельница — женская одежда из льна от производителя в Костроме. Каталог льняной одежды оптом и в розницу, натуральные ткани, доставка и условия сотрудничества.",
            "https://www.kmel.ru/?utm_source=test -> 404.",
            "https://www.kmel.ru/catalog/ -> 302 -> /catalog.",
        ],
    },
    {
        "project": "Витали",
        "url": "https://vitalikostroma.ru/",
        "access": "ICMS-доступ рабочий, вход выполнен. Yandex Webmaster не проверен: в таблице отмечено, что его надо сделать.",
        "summary": "Сильная зона риска — дубли слэш/без слэша и параметров отдают 200 OK, canonical отсутствует. У разделов продуктов одинаковые title, нет description, на страницах встречается по два H1. В админке разделов каталога нет SEO-полей, поэтому безопасной CMS-правки не нашлось.",
        "fixed": [
            "Вошел в ICMS v2.2.",
            "Открыл каталог и редактор раздела 39: есть поля названия, состава, описания, изображений и статуса.",
            "SEO-title, meta description и canonical-полей в редакторе раздела не найдено; правки не выполнялись, чтобы не менять контент вместо шаблона.",
        ],
        "issues": [
            ("P1", "Склеить слэш/без слэша", "Выбрать формат из sitemap и настроить 301 на него для /catalog, /catalog/vitalgar, /products/39 и др."),
            ("P1", "Добавить canonical", "В шаблоне главной, каталога и product-разделов вывести canonical на чистый URL."),
            ("P2", "Уникализировать product SEO", "Для /products/39, /products/30, /products/31 нужны уникальные title и description."),
            ("P2", "Оставить один H1", "В шаблоне product-раздела второй H1 заменить на H2/H3 или заголовок блока."),
        ],
        "evidence": [
            "/products/39 и /products/39/ отдают 200 OK.",
            "/products/39?p=1 отдает дубль с 200 OK.",
            "Редактор раздела 39 не содержит SEO-полей.",
        ],
    },
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(2)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_issue_table(doc, issues):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    widths = [Inches(0.7), Inches(1.8), Inches(4.0)]
    headers = ["Приоритет", "Пункт", "Что сделать"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, "E8EEF5")
        cell.width = widths[i]
    for priority, title, action in issues:
        row = table.add_row()
        values = [priority, title, action]
        for i, value in enumerate(values):
            set_cell_text(row.cells[i], value)
            row.cells[i].width = widths[i]


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"SEO-аудит, {TODAY}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 90, 90)


def build_report(data):
    doc = Document()
    configure_document(doc)
    add_footer(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run(f"SEO-отчет: {data['project']}")
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string("1F4D78")
    run.bold = True

    meta = doc.add_paragraph()
    meta.add_run("Сайт: ").bold = True
    meta.add_run(data["url"])
    meta.add_run(" | Дата проверки: ").bold = True
    meta.add_run(TODAY)

    doc.add_heading("Краткое резюме", level=1)
    doc.add_paragraph(data["summary"])

    doc.add_heading("Доступы и Вебмастер", level=1)
    doc.add_paragraph(data["access"])
    doc.add_paragraph("Важно: favicon не исправлялся и не считался блокером, как указано в постановке задачи.")

    doc.add_heading("Что проверено и исправлено", level=1)
    add_bullets(doc, data["fixed"])

    doc.add_heading("Оставшиеся задачи", level=1)
    add_issue_table(doc, data["issues"])

    doc.add_heading("Доказательства проверки", level=1)
    add_bullets(doc, data["evidence"])

    doc.add_heading("Использованный SEO-чеклист", level=1)
    add_bullets(doc, GUIDELINES)

    doc.add_heading("Следующий шаг", level=1)
    doc.add_paragraph(
        "После правок редиректов и canonical нужно заново проверить сайт в Яндекс.Вебмастере и отправить переобход важных страниц. "
        "Для проектов с недоступным Вебмастером сначала нужен вход в Яндекс ID или выдача доступа на рабочий аккаунт."
    )

    filename = f"SEO отчет - {data['project']}.docx"
    path = OUT_DIR / filename
    doc.save(path)
    return path


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    created = [build_report(report) for report in REPORTS]
    for path in created:
        print(path)


if __name__ == "__main__":
    main()
